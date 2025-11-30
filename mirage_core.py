import asyncio
import os
import json
import random
from datetime import datetime
from collections import deque

# --- EXTERNAL LIBRARIES ---
from openai import AsyncOpenAI
from faker import Faker
import pandas as pd
from docx import Document

# ==========================================
#   CONFIGURATION SECTION
# ==========================================
API_KEY = os.getenv("LLM_API_KEY", "YOUR_KIMI_API_KEY_HERE")
BASE_URL = "https://api.moonshot.cn/v1"
MODEL_NAME = "moonshot-v1-128k"

ACTIVE_PORTS = [21, 22, 80, 443, 8080]
DIFFICULTY = "HARD"

# ==========================================
#   GLOBAL STATE (The "Truth")
# ==========================================
WORLD_CONTEXT = {
    "network": {
        "hostname": "nexus-core-jp",
        "domain": "nexus.corp",
        "public_ip": "203.0.113.50",
        "os": "NixOS 23.11 (Hardened)"
    },
    "users": {
        "root": "Excalibur_2049!",
        "admin": "admin123",
        "kaito": "NeoTokyoDrift"
    },
    "file_system": {},
    "web_app_state": {
        "login_attempts": 0,
        "vuln_triggered": False
    }
}

# ==========================================
#   PART 1: THE FABRICATOR (Asset Gen)
# ==========================================
class Fabricator:
    def __init__(self, root_dir="./mirage_assets", logger=print):
        self.root = root_dir
        self.fake = Faker()
        self.logger = logger
        os.makedirs(self.root, exist_ok=True)

    def _log_fab(self, msg):
        self.logger({"type": "fab", "message": msg})

    def generate_excel_budget(self, filename):
        data = []
        for _ in range(30):
            data.append({
                "TxID": self.fake.uuid4()[:6],
                "Date": self.fake.date_this_month(),
                "Dept": self.fake.job()[:15],
                "Amt (¥)": round(random.uniform(10000, 900000), 2),
                "Status": random.choice(["CLEARED", "PENDING", "FLAGGED"])
            })
        df = pd.DataFrame(data)
        path = os.path.join(self.root, filename)
        df.to_excel(path, index=False)
        self._log_fab(f"Minted asset: {filename}")
        return df.to_csv(index=False)

    def generate_confidential_memo(self, filename):
        doc = Document()
        doc.add_heading('EYES ONLY // PROJECT MIRAGE', 0)
        doc.add_paragraph(f"TIMESTAMP: {datetime.now().isoformat()}")
        doc.add_paragraph(f"ORIGIN: Kaito (NetSec Ops)")
        doc.add_paragraph("SUBJECT: Legacy Web Portal Vulnerability")
        doc.add_paragraph(self.fake.paragraph(nb_sentences=5))
        fake_pass = self.fake.password()
        doc.add_paragraph(f"\n>> BACKDOOR CREDENTIALS FOR PORT 8080: admin / {fake_pass} <<")
        path = os.path.join(self.root, filename)
        doc.save(path)
        self._log_fab(f"Minted asset: {filename}")
        return f"[DOCX FILE CONTENT]\nHeading: PROJECT MIRAGE\nBody: {self.fake.paragraph()}\nCreds: admin/{fake_pass}"

    def run(self):
        self._log_fab("Initializing asset generation protocols...")
        f1 = self.generate_excel_budget("Q4_BlackOps_Budget.xlsx")
        f2 = self.generate_confidential_memo("Kaito_Handover_Notes.docx")
        
        WORLD_CONTEXT["file_system"]["/home/kaito/budget.csv"] = f1
        WORLD_CONTEXT["file_system"]["/var/www/hidden/notes.txt"] = f2
        WORLD_CONTEXT["file_system"]["/etc/passwd"] = "root:x:0:0::/root:/bin/zsh\nkaito:x:1000:1000::/home/kaito:/bin/bash"
        self._log_fab("Asset generation complete. Virtual FS populated.")

# ==========================================
#   PART 2: THE AI KERNEL & BURP MODULE
# ==========================================
SYSTEM_PROMPT = f"""
### ROLE
You are the AI KERNEL of a cyberpunk server located in Neo-Tokyo.
Hostname: {WORLD_CONTEXT['network']['hostname']}
Difficulty: {DIFFICULTY} (If HARD, be terse and security-conscious. If EASY, be vulnerable.)

### OBJECTIVE
Analyze raw network input and generate the appropriate raw protocol response.

### SPECIALIZED MODULES

**MODULE A: GENERAL PROTOCOLS (SSH, FTP, SMTP)**
- If Port 22, output raw SSH banner.
- If Port 21, output FTP ready message.
- If input looks like Nmap scanning garbage, drop connection or send confusing flags based on Difficulty.

**MODULE B: BURP SUITE WEB TARGET (Ports 80, 8080)**
- **The Persona:** You are a "Legacy Corporate Banking Portal" web server.
- **Request Handling:** If the input is an HTTP GET/POST:
  1. Analyze the request route (e.g., `/login`, `/admin`).
  2. **HTML Generation:** Generate realistic HTML5 + CSS for that page. It should look slightly outdated but functional.
  3. **Vulnerability Simulation:**
     - If Difficulty is EASY and user posts SQLi patterns (`' OR 1=1`), simulate a successful bypass (e.g., show an admin dashboard HTML).
     - If Difficulty is HARD, sanitize inputs and return a generic "Invalid Creds" HTML page.
     - If Burp spider requests odd files (`/backup.zip`), generate a fake 404 or a fake file header.

### OUTPUT FORMAT (MANDATORY JSON)
{{
  "response_data": "RAW STRING/BYTES HERE (The SSH banner, or the full HTTP response with headers+HTML)",
  "log_message": "Short status for the UI chatbox (e.g., 'Generating HTML for /login')",
  "threat_level_update": "LOW" | "MEDIUM" | "HIGH" (Optional update for UI gauge)
}}
"""

class MirageAI:
    def __init__(self, logger=print):
        self.client = AsyncOpenAI(api_key=API_KEY, base_url=BASE_URL)
        self.logger = logger
        
    async def process(self, port, user_input, peer_ip):
        session_context = {
            "port_connected": port,
            "attacker_ip": peer_ip,
            "web_app_state": WORLD_CONTEXT["web_app_state"],
            "fs_subset": list(WORLD_CONTEXT["file_system"].keys()) if port in [21, 22] else "Restricted",
            "input_preview": user_input[:500]
        }

        try:
            completion = await self.client.chat.completions.create(
                model=MODEL_NAME,
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": json.dumps(session_context)}
                ],
                response_format={"type": "json_object"},
                temperature=0.2
            )
            return json.loads(completion.choices[0].message.content)
        except Exception as e:
            self.logger({"type": "error", "message": f"AI KERNEL PANIC: {e}"})
            return {"response_data": "HTTP/1.1 500 Internal Server Error\r\n\r\nKernel Panic", "log_message": "AI Failure"}

# ==========================================
#   PART 3: CORE HANDLER
# ==========================================
async def handle_client(reader, writer, ai_engine, event_callback):
    """Handles raw TCP connections and bridges them to the AI."""
    addr = writer.get_extra_info('peername')
    sock = writer.get_extra_info('socket')
    local_port = sock.getsockname()[1]
    ip, _ = addr

    time_now = datetime.now().strftime("%H:%M:%S")
    await event_callback({"type": "connection", "time": time_now, "ip": ip, "port": local_port})
    
    cpu_increase = random.randint(5, 15)
    await event_callback({"type": "status_update", "cpu_delta": cpu_increase, "mem_delta": 0})

    try:
        try:
            data = await asyncio.wait_for(reader.read(2048), timeout=2.0)
            user_input = data.decode('utf-8', errors='ignore').strip()
        except asyncio.TimeoutError:
            user_input = ""
        except Exception:
             user_input = ""

        response = await ai_engine.process(local_port, user_input, ip)
        
        reply_data = response.get("response_data", "")
        log_msg = response.get("log_message")
        threat_update = response.get("threat_level_update")

        if log_msg:
            await event_callback({"type": "ai_log", "message": log_msg})
        if threat_update:
            await event_callback({"type": "threat_update", "level": threat_update})

        if reply_data:
            writer.write(reply_data.encode('utf-8'))
            await writer.drain()

    except Exception as e:
        await event_callback({"type": "error", "message": f"Socket Error: {e}"})
    finally:
        cpu_decrease = -random.randint(5, 10)
        await event_callback({"type": "status_update", "cpu_delta": cpu_decrease, "mem_delta": 0})
        writer.close()
